from docx import Document

doc = Document()
doc.add_heading('Request for Proposal (RFP) - Global Bank Corp', 0)
doc.add_heading('1. Project Requirements', level=1)

doc.add_paragraph('This document outlines the strict requirements for the upcoming enterprise technology modernization project.')

doc.add_paragraph('REQUIREMENT 1: The vendor must demonstrate that their enterprise software system has a guaranteed availability of 99.9% uptime.')
doc.add_paragraph('REQUIREMENT 2: The vendor must provide 24/7/365 global support with response times under 1 hour for critical incidents.')
doc.add_paragraph('REQUIREMENT 3: The solution must natively deploy on and support AWS Azure GCP infrastructures securely.')
doc.add_paragraph('REQUIREMENT 4: The vendor must hold active ISO 27001 and SOC 2 Type II security certifications.')
doc.add_paragraph('REQUIREMENT 5: The vendor must possess distinct AI/ML capabilities, specifically in Document parsing and NLP, demonstrated in similar successful enterprise workflows.')
doc.add_paragraph('REQUIREMENT 6: The overall system architecture must guarantee stringent data protection, including encryption at rest via AES-256 and encryption in transit via TLS 1.3.')

doc.save('demo_rfp.docx')
print('demo_rfp.docx generated successfully!')
