import os
from docx import Document
from PIL import Image, ImageDraw, ImageFont

def make_docx():
    os.makedirs("demo_assets", exist_ok=True)
    doc = Document()
    doc.add_heading("GlobalBank Vendor RFP - 2026", 0)
    
    doc.add_heading("Section 1: Security Requirements", 1)
    doc.add_paragraph("REQ-01: The system must enforce Role-Based Access Control (RBAC) down to the field level.")
    doc.add_paragraph("REQ-02: All customer data in Transit and at Rest must be encrypted using AES-256 and TLS 1.3 standards.")
    doc.add_paragraph("REQ-03: The vendor must provide 24/7 Quantum Threat Remediation in their SLA.") # This one won't be in KB! It will get low confidence and FLAG for review.
    
    doc.add_heading("Section 2: Pricing Requirements", 1)
    doc.add_paragraph("REQ-04: Please confirm the base platform fee for Enterprise deployments.")
    
    doc.save("demo_assets/Golden_Demo_RFP.docx")
    print("Created demo_assets/Golden_Demo_RFP.docx")

def make_img():
    os.makedirs("demo_assets", exist_ok=True)
    img = Image.new('RGB', (800, 300), color=(255, 255, 255))
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 20)
    except:
        font = ImageFont.load_default()
        
    text = (
        "Security Addendum - Scanned Annex\n\n"
        "Requirement: System MUST support SSO.\n"
        "Requirement: System MUST encrypt all data at rest.\n"
        "Requirement: The system shall be SOC2 certified."
    )
    d.text((20, 20), text, fill=(0,0,0), font=font)
    img.save("demo_assets/Scanned_RFP_Table.png")
    print("Created demo_assets/Scanned_RFP_Table.png")

if __name__ == "__main__":
    make_docx()
    make_img()
