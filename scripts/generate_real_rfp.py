import sys
import docx

doc = docx.Document()
doc.add_heading('BidFactory Solutions RFP Requirements', 0)
doc.add_paragraph('REQUIREMENT 1: The vendor must natively support hosting on AWS Azure GCP platforms.') # 1. Technical capabilities
doc.add_paragraph('REQUIREMENT 2: The vendor must possess SOC 2 certification.') # 2. Security/ Compliance
doc.add_paragraph('REQUIREMENT 3: The system must guarantee 99.9% availability for all clients.') # 3. SLA document
doc.add_paragraph('REQUIREMENT 4: The system must feature advanced AI/ML capabilities for MLOps and drift detection.') # 4. AI/ML document
doc.add_paragraph('REQUIREMENT 5: The vendor must have completed previous projects building a Healthcare Data Lake.') # 5. Previous project
doc.add_paragraph('REQUIREMENT 6: The vendor must hold active offices operating in Tokyo, Japan.') # 6. Does NOT exist (Our KB explicitly states SF, NY, London, Singapore).
doc.add_paragraph('REQUIREMENT 7: The vendor must enforce data encryption using AES-256 for all at-rest data, and utilize Quantum key distribution.') # 7. Partial evidence (we AES-256, but NO Quantum keys mentioned in KB)
doc.add_paragraph('REQUIREMENT 8: The vendor must have an RTO (Recovery Time Objective) of 15 minutes.') # 8. Contradiction/Conflict (our KB explicitly states RTO of 1 hour).
doc.save('real_test_rfp.docx')

print("Created real_test_rfp.docx")
