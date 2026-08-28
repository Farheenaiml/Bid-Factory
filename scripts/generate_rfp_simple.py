import sys
import docx

doc = docx.Document()
doc.add_heading('BidFactory RFP')
doc.add_paragraph('This is the RFP text.')
doc.add_paragraph('The vendor must guarantee 99.9% availability.')
doc.add_paragraph('The vendor must hold SOC 2 certification.')
doc.add_paragraph('The vendor must support AWS Azure GCP.')
doc.add_paragraph('The vendor must provide data protection.')
doc.add_paragraph('The vendor must establish ISO 27001.')
doc.add_paragraph('The system must have AI/ML capabilities.')
doc.save('demo_rfp_simple.docx')
