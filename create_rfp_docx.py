from docx import Document

doc = Document()
doc.add_heading('RFP Requirements', 0)

reqs = [
    "The supplier must provide 24/7 support.",
    "The supplier must provide encrypted storage for all data.",
    "The supplier must provide 99.9% availability.",
    "The system must include a certified security module.",
    "The system must support regular automated backups.",
    "The supplier must have a response time of under 1 hour."
]

for r in reqs:
    doc.add_paragraph(r, style='List Bullet')

doc.save('real_rfp.docx')
