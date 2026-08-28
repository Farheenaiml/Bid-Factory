import io
import csv
from uuid import UUID
from docx import Document
from backend.services.repository import repository

class ExportService:
    @staticmethod
    def generate_docx(bid_id: UUID) -> bytes:
        bid = repository.get_bid(bid_id)
        # Gather approved requirements
        items = list(repository._review_items.values())
        bid_items = [r for r in items if r.bid_id == bid_id and r.review_status == "APPROVED"]
        reqs = {r.id: r for r in repository.get_requirements(bid_id)}
        
        doc = Document()
        doc.add_heading(f"Final Bid Response - {bid.rfp.title}", 0)
        
        doc.add_heading("Compliance Matrix & Answers", 1)
        for review in bid_items:
            req_text = reqs[review.requirement_id].requirement_text if review.requirement_id in reqs else str(review.requirement_id)
            doc.add_heading(f"Requirement: {req_text[:50]}...", 2)
            doc.add_paragraph(f"Original Text: {req_text}")
            doc.add_paragraph(f"Compliance Status: {review.compliance_status}")
            doc.add_paragraph(f"Final Answer: {review.proposed_response}")
            doc.add_paragraph("") # Space

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
        
    @staticmethod
    def generate_csv(bid_id: UUID) -> bytes:
        bid = repository.get_bid(bid_id)
        items = list(repository._review_items.values())
        bid_items = [r for r in items if r.bid_id == bid_id and r.review_status == "APPROVED"]
        reqs = {r.id: r for r in repository.get_requirements(bid_id)}
        
        out = io.StringIO()
        writer = csv.writer(out, quoting=csv.QUOTE_MINIMAL)
        writer.writerow(["Requirement ID", "Requirement Text", "Compliance Status", "Final Answer"])
        for review in bid_items:
            req_text = reqs[review.requirement_id].requirement_text if review.requirement_id in reqs else str(review.requirement_id)
            writer.writerow([
                str(review.requirement_id),
                req_text,
                review.compliance_status,
                review.proposed_response
            ])
        return out.getvalue().encode('utf-8')

export_service = ExportService()
