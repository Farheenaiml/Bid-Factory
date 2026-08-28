import hashlib
import io
from pathlib import Path
from typing import Any

from fastapi import HTTPException, status
from pydantic import BaseModel, Field


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}


class ExtractedDocument(BaseModel):
    document_name: str
    document_type: str
    source_path: str
    document_hash: str
    pages: list[dict[str, Any]] = Field(default_factory=list)


class DocumentIngestionService:
    def extract(self, path: Path) -> ExtractedDocument:
        if not path.is_file():
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge-base file was not found.")
        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise HTTPException(status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE, detail="Unsupported file format.")
        contents = path.read_bytes()
        if not contents:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Knowledge-base document is empty.")

        if extension == ".pdf":
            pages = self._extract_pdf(contents)
            document_type = "pdf"
        elif extension in {".png", ".jpg", ".jpeg"}:
            pages = self._extract_image(contents)
            document_type = "image"
        else:
            pages = self._extract_docx(contents)
            document_type = "docx"
        if not any(page["text"].strip() for page in pages):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Knowledge-base document contains no extractable text.")
        return ExtractedDocument(
            document_name=path.name,
            document_type=document_type,
            source_path=str(path.resolve()),
            document_hash=hashlib.sha256(contents).hexdigest(),
            pages=pages,
        )

    @staticmethod
    def _extract_pdf(contents: bytes) -> list[dict[str, Any]]:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(contents))
            return [{"text": page.extract_text() or "", "page_number": index + 1, "section": None} for index, page in enumerate(reader.pages)]
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=422, detail="Unable to extract text from PDF document.") from exc

    @staticmethod
    def _extract_image(contents: bytes) -> list[dict[str, Any]]:
        try:
            from PIL import Image
            import pytesseract
            import sys
            
            if sys.platform == 'win32':
                pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

            image = Image.open(io.BytesIO(contents))
            # Extract real text dynamically
            text = pytesseract.image_to_string(image)
                
            return [{"text": text or "", "page_number": 1, "section": None}]
        except Exception as exc:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=422, detail="Unable to extract text from Image using real OCR.") from exc

    @staticmethod
    def _extract_docx(contents: bytes) -> list[dict[str, Any]]:
        try:
            from docx import Document

            document = Document(io.BytesIO(contents))
            section: str | None = None
            paragraphs: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if paragraph.style.name.lower().startswith("heading"):
                    section = text
                paragraphs.append(text)
            return [{"text": "\n".join(paragraphs), "page_number": None, "section": section}]
        except Exception as exc:
            raise HTTPException(status_code=422, detail="Unable to extract text from DOCX document.") from exc


document_ingestion_service = DocumentIngestionService()