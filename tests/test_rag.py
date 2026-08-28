from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from backend.services.chunking import ChunkingService
from backend.services.document_ingestion import DocumentIngestionService
from backend.services.knowledge_base import KnowledgeBaseService


class KeywordEmbeddings:
    def embed(self, texts: list[str]) -> list[list[float]]:
        return [
            [
                float("security" in text.lower()),
                float("pricing" in text.lower()),
                float("support" in text.lower()),
            ]
            for text in texts
        ]


def create_docx(path: Path, text: str, heading: str = "Policy") -> None:
    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(text)
    document.save(path)


def create_pdf(path: Path, text: str) -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
    })
    page[NameObject("/Resources")] = DictionaryObject({
        NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})
    })
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode())
    page[NameObject("/Contents")] = stream
    writer.write(str(path))


def build_service(tmp_path: Path) -> KnowledgeBaseService:
    knowledge_base = tmp_path / "knowledge_base"
    vector_store = tmp_path / "vector_store"
    knowledge_base.mkdir()
    return KnowledgeBaseService(
        knowledge_base_dir=knowledge_base,
        vector_store_dir=vector_store,
        embedding_model="test-only",
        top_k=3,
        min_score=0.5,
        embeddings=KeywordEmbeddings(),
    )


def test_docx_ingestion_preserves_metadata_and_creates_chunks(tmp_path: Path) -> None:
    path = tmp_path / "security.docx"
    create_docx(path, "Our security policy requires annual reviews.")

    document = DocumentIngestionService().extract(path)
    chunks = ChunkingService(chunk_size=40, overlap=5).create_chunks(document)

    assert document.document_name == "security.docx"
    assert document.document_type == "docx"
    assert document.source_path == str(path.resolve())
    assert document.pages[0]["section"] == "Policy"
    assert chunks
    assert chunks[0].document_hash == document.document_hash
    assert chunks[0].section == "Policy"


def test_pdf_ingestion_extracts_text_and_page_metadata(tmp_path: Path) -> None:
    path = tmp_path / "security.pdf"
    create_pdf(path, "Security policy")

    document = DocumentIngestionService().extract(path)

    assert document.document_type == "pdf"
    assert document.pages[0]["page_number"] == 1
    assert "Security policy" in document.pages[0]["text"]


def test_retrieval_returns_relevant_evidence_and_metadata(tmp_path: Path) -> None:
    service = build_service(tmp_path)
    path = tmp_path / "security.docx"
    create_docx(path, "Our security policy requires annual reviews.", "Security")
    service.ingest_file(path)

    response = service.search("security requirements")

    assert response["message"] == "evidence found"
    assert len(response["results"]) == 1
    result = response["results"][0]
    assert result.document_name == "security.docx"
    assert result.source_path == str(path.resolve())
    assert result.section == "Security"
    assert result.similarity_score >= 0.7
    assert "hybrid_scores" in result.metadata
    assert result.metadata["document_type"] == "docx"


def test_low_relevance_query_returns_no_evidence(tmp_path: Path) -> None:
    service = build_service(tmp_path)
    path = tmp_path / "security.docx"
    create_docx(path, "Our security policy requires annual reviews.")
    service.ingest_file(path)

    response = service.search("pricing requirements")

    assert response["results"] == []
    assert response["message"] == "no relevant evidence found"


def test_duplicate_document_is_not_indexed_twice(tmp_path: Path) -> None:
    service = build_service(tmp_path)
    path = tmp_path / "security.docx"
    create_docx(path, "Our security policy requires annual reviews.")

    first = service.ingest_file(path)
    duplicate = service.ingest_file(path)
    response = service.search("security requirements")

    assert first.chunk_count > 0
    assert duplicate.document_hash == first.document_hash
    assert len(response["results"]) == 1
